"""
Export VaultScan security reports as PDF and DOCX.

Mirrors the on-screen report: executive brief, score, severity/service charts,
full finding cards (what / why / fix / compliance), priority plan, compliance,
technical notes, glossary.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.graphics.charts.barcharts import HorizontalBarChart
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.shapes import Circle, Drawing, Rect, String
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
    Flowable,
)

SEV_COLORS = {
    "CRITICAL": colors.HexColor("#ff3d57"),
    "HIGH": colors.HexColor("#ff9900"),
    "MEDIUM": colors.HexColor("#3874ff"),
    "LOW": colors.HexColor("#6b7280"),
}
SEV_LABELS = {
    "CRITICAL": "Critical — fix immediately",
    "HIGH": "High — fix this week",
    "MEDIUM": "Medium — plan remediation",
    "LOW": "Low — harden when possible",
}
GLOSSARY = [
    (
        "Posture score",
        "0–100 health score. 100 means no known misconfigurations in this scan.",
    ),
    (
        "Critical",
        "Actively dangerous exposure (e.g. public data or world-open admin ports).",
    ),
    (
        "High",
        "Serious weakness that attackers often chain with other issues.",
    ),
    (
        "Medium / Low",
        "Hardening and compliance gaps; less urgent but still important.",
    ),
    (
        "CIS / NIST / GDPR",
        "Industry and legal frameworks. Findings map to rules auditors care about.",
    ),
]


def _utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _safe(text: Any, limit: int = 8000) -> str:
    s = str(text or "").replace("\x00", "")
    if len(s) > limit:
        return s[: limit - 1] + "…"
    return s


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def why_it_matters(finding: dict[str, Any]) -> str:
    """Same plain-language logic as the web report."""
    if finding.get("why_it_matters"):
        return str(finding["why_it_matters"])
    sev = str(finding.get("severity") or "").upper()
    service = str(finding.get("service") or "").upper()
    title = str(finding.get("title") or finding.get("description") or "").lower()

    if "public" in title or "0.0.0.0" in title or "world" in title:
        return "Anyone on the internet may reach this resource without being your employee or customer."
    if "encrypt" in title:
        return "If disks or objects are stolen or copied, data may be readable without a key."
    if "mfa" in title:
        return "A stolen password alone could let an attacker sign in as this identity."
    if "admin" in title or "permissive" in title or "*" in title:
        return "If this identity is compromised, the attacker may control large parts of the account."
    if service == "S3":
        return "Storage misconfigurations are a leading cause of cloud data breaches."
    if service == "EC2":
        return "Network exposure on compute often leads to intrusion and lateral movement."
    if service == "IAM":
        return "Identity weaknesses turn a small compromise into full account takeover."
    if service == "RDS":
        return "Database exposure can leak customer or business-critical records."
    if sev == "CRITICAL":
        return "Highest priority — treat as an active security incident until fixed."
    if sev == "HIGH":
        return "Should be fixed soon; increases risk of breach or audit failure."
    return "Improves resilience and compliance; schedule alongside other hardening work."


def build_report_context(
    scan: dict[str, Any],
    narrative: dict[str, str] | None = None,
) -> dict[str, Any]:
    summary = scan.get("summary") or {}
    raw = scan.get("vulnerabilities") or scan.get("findings") or []
    score = int(scan.get("score") or 0)
    if score >= 90:
        risk = "LOW"
    elif score >= 70:
        risk = "MODERATE"
    elif score >= 40:
        risk = "HIGH"
    else:
        risk = "CRITICAL"

    findings: list[dict[str, Any]] = []
    service_counts: dict[str, int] = {}
    for f in raw:
        sev = str(f.get("severity") or "LOW").upper()
        svc = str(f.get("service") or "OTHER")
        service_counts[svc] = service_counts.get(svc, 0) + 1
        findings.append(
            {
                "severity": sev,
                "service": svc,
                "resource": f.get("id") or f.get("resource") or "—",
                "title": f.get("title") or "",
                "description": f.get("description") or "",
                "remediation": f.get("remediation") or "",
                "compliance": f.get("compliance") or [],
                "why_it_matters": why_it_matters(f),
                "region": f.get("region") or "",
            }
        )

    # Severity order for display
    order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    findings.sort(key=lambda x: order.get(x["severity"], 9))

    by_service = [
        {"service": k, "count": c}
        for k, c in sorted(service_counts.items(), key=lambda x: -x[1])
    ]
    by_severity = [
        {
            "severity": sev,
            "count": int(summary.get(sev, 0)),
            "label": SEV_LABELS[sev],
        }
        for sev in ("CRITICAL", "HIGH", "MEDIUM", "LOW")
    ]

    narr = narrative or {}
    return {
        "scan_id": scan.get("scan_id") or "SCAN",
        "timestamp": scan.get("timestamp") or _utc_now(),
        "account_id": scan.get("account_id") or "—",
        "region": scan.get("region") or "—",
        "mode": scan.get("mode") or "—",
        "role_arn": scan.get("role_arn") or "",
        "score": score,
        "risk": (narr.get("risk_level") or risk).upper(),
        "total": scan.get("total_findings") or len(findings),
        "summary": {
            "CRITICAL": int(summary.get("CRITICAL", 0)),
            "HIGH": int(summary.get("HIGH", 0)),
            "MEDIUM": int(summary.get("MEDIUM", 0)),
            "LOW": int(summary.get("LOW", 0)),
        },
        "headline": narr.get("headline")
        or f"Cloud security posture is {risk.lower()} ({score}/100)",
        "executive_summary": narr.get("executive_summary")
        or f"VaultScan found {len(findings)} issue(s) in this environment.",
        "what_this_means": narr.get("what_this_means")
        or "Misconfigurations increase breach and audit risk.",
        "priority_actions": narr.get("priority_actions") or "Review findings below.",
        "technical_notes": narr.get("technical_notes")
        or "Re-scan after remediation to verify score improvement.",
        "findings": findings,
        "by_severity": by_severity,
        "by_service": by_service,
        "compliance": scan.get("compliance") or [],
        "glossary": GLOSSARY,
    }


# ─── PDF graphics ─────────────────────────────────────────────────────────────

def _score_drawing(score: int, size: float = 120) -> Drawing:
    d = Drawing(size, size)
    cx, cy, r = size / 2, size / 2, size * 0.38
    if score >= 90:
        col = colors.HexColor("#00a854")
    elif score >= 70:
        col = colors.HexColor("#ff9900")
    else:
        col = colors.HexColor("#ff3d57")

    d.add(
        Circle(
            cx,
            cy,
            r + 6,
            fillColor=None,
            strokeColor=colors.HexColor("#e5e7eb"),
            strokeWidth=8,
        )
    )
    # Approximate arc with thick circle segment is hard; use filled ring + overlay text
    d.add(
        Circle(
            cx,
            cy,
            r + 6,
            fillColor=None,
            strokeColor=col,
            strokeWidth=8,
            strokeOpacity=min(1.0, 0.35 + score / 100 * 0.65),
        )
    )
    d.add(
        String(
            cx,
            cy + 4,
            str(score),
            fontSize=22,
            fillColor=col,
            textAnchor="middle",
            fontName="Helvetica-Bold",
        )
    )
    d.add(
        String(
            cx,
            cy - 14,
            "/ 100",
            fontSize=8,
            fillColor=colors.HexColor("#6b7280"),
            textAnchor="middle",
            fontName="Helvetica",
        )
    )
    return d


def _severity_pie(by_severity: list[dict]) -> Drawing | None:
    data = [int(x["count"]) for x in by_severity if int(x["count"]) > 0]
    labels = [x["severity"] for x in by_severity if int(x["count"]) > 0]
    if not data:
        return None
    d = Drawing(220, 160)
    pie = Pie()
    pie.x = 30
    pie.y = 15
    pie.width = 120
    pie.height = 120
    pie.data = data
    pie.labels = labels
    pie.sideLabels = True
    pie.simpleLabels = False
    pie.slices.strokeWidth = 0.5
    pie.slices.strokeColor = colors.white
    palette = [SEV_COLORS.get(l, colors.grey) for l in labels]
    for i, c in enumerate(palette):
        pie.slices[i].fillColor = c
    d.add(pie)
    return d


def _service_bars(by_service: list[dict]) -> Drawing | None:
    if not by_service:
        return None
    items = by_service[:8]
    n = len(items)
    height = max(100, 28 * n + 40)
    d = Drawing(260, height)
    chart = HorizontalBarChart()
    chart.x = 70
    chart.y = 15
    chart.height = height - 35
    chart.width = 170
    chart.data = [[int(x["count"]) for x in items]]
    chart.categoryAxis.categoryNames = [str(x["service"]) for x in items]
    chart.bars[0].fillColor = colors.HexColor("#3874ff")
    chart.bars[0].strokeColor = None
    chart.valueAxis.valueMin = 0
    chart.valueAxis.valueMax = max(int(x["count"]) for x in items) + 1
    chart.valueAxis.valueStep = 1
    chart.categoryAxis.labels.fontSize = 8
    chart.valueAxis.labels.fontSize = 7
    d.add(chart)
    return d


def _severity_legend_table(by_severity: list[dict], styles) -> Table:
    rows = []
    for item in by_severity:
        sev = item["severity"]
        count = int(item["count"])
        label = item.get("label") or SEV_LABELS.get(sev, sev)
        col = SEV_COLORS.get(sev, colors.grey)
        # Color swatch as a small table cell background
        rows.append(
            [
                Paragraph(f"<b>{_escape(sev)}</b>", styles["VSSmall"]),
                Paragraph(str(count), styles["VSSmall"]),
                Paragraph(_escape(label), styles["VSSmall"]),
            ]
        )
    t = Table(rows, colWidths=[28 * mm, 14 * mm, 70 * mm])
    style_cmds = [
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8f9fb")),
        ("BOX", (0, 0), (-1, -1), 0.3, colors.HexColor("#e5e7eb")),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, colors.HexColor("#e5e7eb")),
    ]
    for i, item in enumerate(by_severity):
        c = SEV_COLORS.get(item["severity"], colors.grey)
        style_cmds.append(("TEXTCOLOR", (0, i), (0, i), c))
        style_cmds.append(("FONTNAME", (0, i), (0, i), "Helvetica-Bold"))
    t.setStyle(TableStyle(style_cmds))
    return t


class FindingCard(Flowable):
    """Full-detail finding block matching the web report row content."""

    def __init__(self, index: int, finding: dict[str, Any], width: float, styles):
        super().__init__()
        self.index = index
        self.f = finding
        self.box_width = width
        self.styles = styles
        self._inner: Table | None = None
        self._height = 0

    def wrap(self, availWidth, availHeight):
        w = min(self.box_width, availWidth)
        f = self.f
        sev = f.get("severity") or "LOW"
        title = f.get("title") or "Finding"
        res = f.get("resource") or "—"
        desc = f.get("description") or ""
        why = f.get("why_it_matters") or ""
        fix = f.get("remediation") or "See remediation hub in VaultScan."
        comp = f.get("compliance") or []
        comp_str = " · ".join(str(c) for c in comp[:6]) if comp else "—"

        header = (
            f"<b>#{self.index}  [{_escape(sev)}]</b>  {_escape(str(f.get('service') or ''))}"
            f"  ·  <font color='#6b7280'>{_escape(_safe(res, 80))}</font>"
        )
        body = (
            f"<b>What we found</b><br/>"
            f"<b>{_escape(_safe(title, 200))}</b><br/>"
            f"{_escape(_safe(desc, 900))}<br/><br/>"
            f"<b>Why it matters</b><br/>"
            f"{_escape(_safe(why, 500))}<br/><br/>"
            f"<b>How to fix</b><br/>"
            f"<font face='Courier' size='7'>{_escape(_safe(fix, 700))}</font><br/><br/>"
            f"<b>Compliance</b><br/>"
            f"<font size='7' color='#3874ff'>{_escape(comp_str)}</font>"
        )
        inner = Table(
            [
                [Paragraph(header, self.styles["VSCardHead"])],
                [Paragraph(body, self.styles["VSCardBody"])],
            ],
            colWidths=[w - 4],
        )
        border = SEV_COLORS.get(sev, colors.HexColor("#6b7280"))
        inner.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111217")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#fafbfc")),
                    ("BOX", (0, 0), (-1, -1), 1, border),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        self._inner = inner
        iw, ih = inner.wrap(w, availHeight)
        self._height = ih
        self.width = w
        self.height = ih
        return w, ih

    def draw(self):
        if self._inner:
            self._inner.drawOn(self.canv, 0, 0)


def _styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="VSTitle",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=colors.HexColor("#0b0c10"),
            spaceAfter=4,
            fontName="Helvetica-Bold",
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSSub",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#555555"),
            spaceAfter=8,
            leading=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSH2",
            parent=styles["Heading2"],
            fontSize=11,
            textColor=colors.HexColor("#111217"),
            spaceBefore=14,
            spaceAfter=4,
            fontName="Helvetica-Bold",
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSHint",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.HexColor("#6b7280"),
            spaceAfter=8,
            leading=11,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSSmall",
            parent=styles["Normal"],
            fontSize=8,
            leading=11,
            textColor=colors.HexColor("#333333"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSCenter",
            parent=styles["Normal"],
            fontSize=8,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#666666"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSCardHead",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.white,
            leading=12,
            fontName="Helvetica-Bold",
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSCardBody",
            parent=styles["Normal"],
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor("#1f2937"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="VSBadge",
            parent=styles["Normal"],
            fontSize=9,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#ff3d57"),
        )
    )
    return styles


# ─── PDF ──────────────────────────────────────────────────────────────────────

def export_pdf(ctx: dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    page_w, page_h = A4
    margin = 14 * mm
    content_w = page_w - 2 * margin

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=f"VaultScan Report {ctx['scan_id']}",
        author="VaultScan Cloud Assistant",
    )
    styles = _styles()
    story: list[Any] = []

    # ── Header (matches web masthead)
    story.append(Paragraph("CLOUD SECURITY POSTURE REPORT", styles["VSTitle"]))
    story.append(
        Paragraph(
            f"<b>{_escape(_safe(ctx['headline'], 300))}</b>",
            styles["VSBody"],
        )
    )
    risk = _safe(ctx["risk"], 20)
    story.append(
        Paragraph(
            f"Scan <b>{_escape(_safe(ctx['scan_id'], 40))}</b>  ·  "
            f"Risk <b>{_escape(risk)}</b>  ·  "
            f"Score <b>{ctx['score']}/100</b>  ·  "
            f"{ctx['total']} finding(s)<br/>"
            f"Account {_escape(_safe(ctx['account_id'], 40))}  ·  "
            f"Region {_escape(_safe(ctx['region'], 20))}  ·  "
            f"Mode {_escape(_safe(ctx['mode'], 20))}<br/>"
            f"Generated with Cloud Assistant · {_utc_now()} · Confidential",
            styles["VSSub"],
        )
    )
    if ctx.get("role_arn"):
        story.append(
            Paragraph(
                f"Role: <font face='Courier' size='8'>{_escape(_safe(ctx['role_arn'], 120))}</font>",
                styles["VSHint"],
            )
        )
    story.append(
        HRFlowable(
            width="100%", thickness=1.5, color=colors.HexColor("#3874ff"), spaceAfter=10
        )
    )

    # ── Score + executive (2-col like web)
    score_draw = _score_drawing(int(ctx["score"]), 110)
    left_cell = [
        Paragraph("POSTURE SCORE", styles["VSH2"]),
        Paragraph("Higher is safer.", styles["VSHint"]),
        score_draw,
        Paragraph(
            f"<para alignment='center'><b>{ctx['total']}</b> total finding(s)</para>",
            styles["VSSmall"],
        ),
    ]
    # Wrap left column content
    left_tbl = Table(
        [[left_cell[0]], [left_cell[1]], [left_cell[2]], [left_cell[3]]],
        colWidths=[58 * mm],
    )
    left_tbl.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 2), (0, 2), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f4f6fb")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    right_inner = Table(
        [
            [Paragraph("EXECUTIVE BRIEF", styles["VSH2"])],
            [
                Paragraph(
                    "Written for managers — what matters in plain language.",
                    styles["VSHint"],
                )
            ],
            [Paragraph(_escape(_safe(ctx["executive_summary"])), styles["VSBody"])],
            [
                Paragraph(
                    "<b>WHAT THIS MEANS FOR THE BUSINESS</b>",
                    styles["VSSmall"],
                )
            ],
            [Paragraph(_escape(_safe(ctx["what_this_means"])), styles["VSBody"])],
        ],
        colWidths=[content_w - 66 * mm],
    )
    right_inner.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f4f6fb")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ("BACKGROUND", (0, 3), (-1, 4), colors.HexColor("#eef2ff")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )

    top_row = Table([[left_tbl, right_inner]], colWidths=[62 * mm, content_w - 64 * mm])
    top_row.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (0, 0), 6),
                ("LEFTPADDING", (1, 0), (1, 0), 6),
            ]
        )
    )
    story.append(top_row)
    story.append(Spacer(1, 10))

    # ── Metrics strip
    s = ctx["summary"]
    snap_data = [
        ["Score", "Total", "Critical", "High", "Medium", "Low"],
        [
            str(ctx["score"]),
            str(ctx["total"]),
            str(s.get("CRITICAL", 0)),
            str(s.get("HIGH", 0)),
            str(s.get("MEDIUM", 0)),
            str(s.get("LOW", 0)),
        ],
    ]
    snap = Table(snap_data, colWidths=[content_w / 6.0] * 6)
    snap.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111217")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("BACKGROUND", (0, 1), (-1, 1), colors.white),
                ("TEXTCOLOR", (2, 1), (2, 1), colors.HexColor("#ff3d57")),
                ("TEXTCOLOR", (3, 1), (3, 1), colors.HexColor("#ff9900")),
                ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d1d5db")),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(snap)

    # ── Charts row
    story.append(Paragraph("FINDINGS BY SEVERITY", styles["VSH2"]))
    story.append(
        Paragraph(
            "How serious are the issues? Critical and High should be fixed first.",
            styles["VSHint"],
        )
    )
    pie = _severity_pie(ctx["by_severity"])
    legend = _severity_legend_table(ctx["by_severity"], styles)
    if pie:
        chart_row = Table([[pie, legend]], colWidths=[100 * mm, content_w - 102 * mm])
        chart_row.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.append(chart_row)
    else:
        story.append(legend)

    story.append(Paragraph("FINDINGS BY SERVICE", styles["VSH2"]))
    story.append(
        Paragraph("Which AWS services have the most problems?", styles["VSHint"])
    )
    bars = _service_bars(ctx["by_service"])
    if bars:
        story.append(bars)
    else:
        story.append(Paragraph("No service breakdown.", styles["VSBody"]))

    # ── Priority plan
    story.append(Paragraph("PRIORITY ACTION PLAN", styles["VSH2"]))
    story.append(
        Paragraph("Start at the top. Complete these, then re-scan.", styles["VSHint"])
    )
    for line in _safe(ctx["priority_actions"], 5000).splitlines():
        line = line.strip()
        if line:
            story.append(Paragraph(_escape(line), styles["VSBody"]))

    # ── Detailed findings (full cards — same depth as web)
    story.append(Paragraph("DETAILED FINDINGS", styles["VSH2"]))
    story.append(
        Paragraph(
            "Each block is one problem: what we found, why it matters, how to fix, and compliance tags.",
            styles["VSHint"],
        )
    )
    if not ctx["findings"]:
        story.append(
            Paragraph(
                "No findings — nothing to remediate in this scan.", styles["VSBody"]
            )
        )
    else:
        for i, f in enumerate(ctx["findings"], 1):
            story.append(Spacer(1, 6))
            story.append(FindingCard(i, f, content_w, styles))

    # ── Compliance
    if ctx["compliance"]:
        story.append(Paragraph("COMPLIANCE SNAPSHOT", styles["VSH2"]))
        story.append(
            Paragraph(
                "How this environment looks against common audit frameworks.",
                styles["VSHint"],
            )
        )
        crow = [
            [
                Paragraph("<b>Framework</b>", styles["VSSmall"]),
                Paragraph("<b>Version</b>", styles["VSSmall"]),
                Paragraph("<b>Status</b>", styles["VSSmall"]),
                Paragraph("<b>Passed</b>", styles["VSSmall"]),
                Paragraph("<b>Coverage</b>", styles["VSSmall"]),
            ]
        ]
        for fw in ctx["compliance"]:
            total = int(fw.get("total") or 0) or 1
            passed = int(fw.get("passed") or 0)
            pct = round(passed / total * 100)
            controls = fw.get("controls") or []
            ctrl_txt = ""
            if controls:
                ctrl_txt = "<br/><font color='#6b7280' size='7'>" + _escape(
                    _safe(
                        " · ".join(
                            f"{c.get('label', '')}: {c.get('value', 0)}%"
                            for c in controls
                        ),
                        220,
                    )
                ) + "</font>"
            crow.append(
                [
                    Paragraph(
                        _escape(_safe(fw.get("name"), 80)) + ctrl_txt,
                        styles["VSSmall"],
                    ),
                    Paragraph(_escape(_safe(fw.get("version"), 24)), styles["VSSmall"]),
                    Paragraph(_escape(_safe(fw.get("status"), 16)), styles["VSSmall"]),
                    Paragraph(f"{passed} / {fw.get('total', 0)}", styles["VSSmall"]),
                    Paragraph(f"{pct}%", styles["VSSmall"]),
                ]
            )
        ct = Table(
            crow,
            colWidths=[
                content_w * 0.38,
                content_w * 0.14,
                content_w * 0.14,
                content_w * 0.16,
                content_w * 0.14,
            ],
        )
        ct.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111217")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#e5e7eb")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#f9fafb")],
                    ),
                ]
            )
        )
        story.append(ct)

    # ── Technical notes + glossary (2-col)
    story.append(Paragraph("TECHNICAL NOTES", styles["VSH2"]))
    story.append(
        Paragraph("For security engineers and DevOps.", styles["VSHint"])
    )
    story.append(Paragraph(_escape(_safe(ctx["technical_notes"])), styles["VSBody"]))

    story.append(Paragraph("GLOSSARY", styles["VSH2"]))
    story.append(
        Paragraph(
            "Quick definitions so the whole team can read this report.",
            styles["VSHint"],
        )
    )
    for term, meaning in ctx.get("glossary") or GLOSSARY:
        story.append(
            Paragraph(
                f"<b>{_escape(term)}</b> — {_escape(meaning)}",
                styles["VSBody"],
            )
        )

    story.append(Spacer(1, 16))
    story.append(
        HRFlowable(
            width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=6
        )
    )
    story.append(
        Paragraph(
            "V.A.U.L.T.S.C.A.N. · CLOUD ASSISTANT REPORT · CONFIDENTIAL",
            styles["VSCenter"],
        )
    )

    doc.build(story)
    return buf.getvalue()


# ─── DOCX (same depth as PDF/web) ─────────────────────────────────────────────

def export_docx(ctx: dict[str, Any]) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading("Cloud Security Posture Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run(_safe(ctx["headline"]))
    run.bold = True
    run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Scan ID: {ctx['scan_id']}  ·  Risk: {ctx['risk']}  ·  Score: {ctx['score']}/100  ·  "
        f"{ctx['total']} findings\n"
        f"Account: {ctx['account_id']}  ·  Region: {ctx['region']}  ·  Mode: {ctx['mode']}\n"
        f"Generated with Cloud Assistant · {_utc_now()} · Confidential"
    ).font.size = Pt(9)

    doc.add_heading("Executive brief", level=1)
    doc.add_paragraph(_safe(ctx["executive_summary"]))

    doc.add_heading("What this means for the business", level=1)
    doc.add_paragraph(_safe(ctx["what_this_means"]))

    doc.add_heading("Posture snapshot", level=1)
    snap = doc.add_table(rows=2, cols=6)
    snap.style = "Table Grid"
    headers = ["Score", "Total", "Critical", "High", "Medium", "Low"]
    values = [
        str(ctx["score"]),
        str(ctx["total"]),
        str(ctx["summary"].get("CRITICAL", 0)),
        str(ctx["summary"].get("HIGH", 0)),
        str(ctx["summary"].get("MEDIUM", 0)),
        str(ctx["summary"].get("LOW", 0)),
    ]
    for i, h in enumerate(headers):
        snap.rows[0].cells[i].text = h
        snap.rows[1].cells[i].text = values[i]

    doc.add_heading("Findings by severity", level=1)
    for item in ctx["by_severity"]:
        doc.add_paragraph(
            f"{item['severity']}: {item['count']} — {item.get('label', '')}",
            style="List Bullet",
        )

    doc.add_heading("Findings by service", level=1)
    for item in ctx["by_service"]:
        doc.add_paragraph(
            f"{item['service']}: {item['count']}",
            style="List Bullet",
        )

    doc.add_heading("Priority action plan", level=1)
    for line in _safe(ctx["priority_actions"], 5000).splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(
                line,
                style="List Number" if line[:1].isdigit() else "List Bullet",
            )

    doc.add_heading("Detailed findings", level=1)
    if not ctx["findings"]:
        doc.add_paragraph("No findings — clean posture for this scan.")
    else:
        for i, f in enumerate(ctx["findings"], 1):
            h = doc.add_heading(
                f"#{i} [{f.get('severity')}] {f.get('service')} — {f.get('resource')}",
                level=2,
            )
            doc.add_paragraph().add_run("What we found").bold = True
            if f.get("title"):
                tp = doc.add_paragraph()
                r = tp.add_run(_safe(f["title"]))
                r.bold = True
            doc.add_paragraph(_safe(f.get("description")))

            doc.add_paragraph().add_run("Why it matters").bold = True
            doc.add_paragraph(_safe(f.get("why_it_matters")))

            doc.add_paragraph().add_run("How to fix").bold = True
            fix = doc.add_paragraph(_safe(f.get("remediation") or "See VaultScan UI"))
            for run in fix.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8)

            comp = f.get("compliance") or []
            if comp:
                doc.add_paragraph().add_run("Compliance").bold = True
                doc.add_paragraph(" · ".join(str(c) for c in comp))

    if ctx["compliance"]:
        doc.add_heading("Compliance snapshot", level=1)
        ct = doc.add_table(rows=1, cols=5)
        ct.style = "Table Grid"
        for i, name in enumerate(
            ["Framework", "Version", "Status", "Passed / Total", "Coverage"]
        ):
            ct.rows[0].cells[i].text = name
        for fw in ctx["compliance"]:
            total = int(fw.get("total") or 0) or 1
            passed = int(fw.get("passed") or 0)
            pct = round(passed / total * 100)
            row = ct.add_row().cells
            row[0].text = _safe(fw.get("name"), 80)
            row[1].text = _safe(fw.get("version"), 30)
            row[2].text = _safe(fw.get("status"), 20)
            row[3].text = f"{passed} / {fw.get('total', 0)}"
            row[4].text = f"{pct}%"
            for c in fw.get("controls") or []:
                doc.add_paragraph(
                    f"  · {c.get('label')}: {c.get('value')}%",
                    style="List Bullet",
                )

    doc.add_heading("Technical notes", level=1)
    doc.add_paragraph(_safe(ctx["technical_notes"]))

    doc.add_heading("Glossary", level=1)
    for term, meaning in ctx.get("glossary") or GLOSSARY:
        p = doc.add_paragraph()
        r = p.add_run(f"{term} — ")
        r.bold = True
        p.add_run(meaning)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run(
        "V.A.U.L.T.S.C.A.N. · CLOUD ASSISTANT REPORT · CONFIDENTIAL"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Fix change report (remediation) PDF / DOCX ───────────────────────────────

def export_fix_report_pdf(report: dict[str, Any]) -> bytes:
    """Professional PDF: before → changed → after, CLI, AI summary."""
    buf = io.BytesIO()
    margin = 14 * mm
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=f"VaultScan Fix Report {report.get('job_id') or ''}",
        author="VaultScan Cloud Assistant",
    )
    styles = _styles()
    story: list[Any] = []
    counts = report.get("counts") or {}
    job_id = _safe(report.get("job_id"), 40)
    delta = report.get("score_delta")
    delta_s = "—" if delta is None else f"{delta:+}"

    story.append(Paragraph("REMEDIATION CHANGE REPORT", styles["VSTitle"]))
    story.append(
        Paragraph(
            "How the environment looked <b>before</b>, what VaultScan <b>changed</b>, "
            "and the expected state <b>after</b> — with CLI and Cloud Assistant notes.",
            styles["VSBody"],
        )
    )
    story.append(
        Paragraph(
            f"Report <b>{_escape(_safe(report.get('report_id'), 48))}</b>  ·  "
            f"Job <b>{_escape(job_id)}</b>  ·  Status {_escape(_safe(report.get('job_status'), 24))}<br/>"
            f"Score before <b>{_escape(str(report.get('score_before') if report.get('score_before') is not None else '—'))}</b>  →  "
            f"after <b>{_escape(str(report.get('score_after') if report.get('score_after') is not None else '—'))}</b>  "
            f"(Δ {_escape(delta_s)})<br/>"
            f"Applied {counts.get('applied', 0)}  ·  Failed {counts.get('failed', 0)}  ·  "
            f"Skipped {counts.get('skipped', 0)}  ·  Restored {counts.get('rolled_back', 0)}  ·  "
            f"Total {counts.get('total', 0)}<br/>"
            f"Generated {_utc_now()}"
            + (" · Cloud Assistant" if report.get("ai_used") else ""),
            styles["VSSub"],
        )
    )
    story.append(
        HRFlowable(
            width="100%", thickness=1.5, color=colors.HexColor("#3874ff"), spaceAfter=10
        )
    )

    story.append(Paragraph("1. Executive summary", styles["VSH2"]))
    story.append(
        Paragraph(_escape(_safe(report.get("executive_summary"), 4000)), styles["VSBody"])
    )
    story.append(Spacer(1, 6))

    recs = report.get("recommendations") or []
    if recs:
        story.append(Paragraph("2. Recommendations", styles["VSH2"]))
        for r in recs:
            story.append(Paragraph(f"• {_escape(_safe(r, 500))}", styles["VSBody"]))
        story.append(Spacer(1, 6))

    story.append(Paragraph("3. Detailed changes (before → after)", styles["VSH2"]))
    story.append(
        Paragraph(
            "Each item shows the pre-change snapshot, what the job did, expected result, "
            "and paste-ready AWS CLI.",
            styles["VSHint"],
        )
    )
    story.append(Spacer(1, 4))

    for i, ch in enumerate(report.get("changes") or [], 1):
        rid = _safe(ch.get("rule_id"), 40)
        res = _safe(ch.get("resource"), 80)
        status = _safe(ch.get("status"), 24)
        risk = _safe(ch.get("risk"), 16)
        block: list[Any] = [
            Paragraph(
                f"<b>{i}. {_escape(rid)}</b>  ·  {_escape(status)}  ·  risk {_escape(risk)}",
                styles["VSH2"],
            ),
            Paragraph(
                f"Resource: <font face='Courier' size='8'>{_escape(res)}</font>",
                styles["VSHint"],
            ),
            Paragraph(f"<b>{_escape(_safe(ch.get('title') or ch.get('summary'), 200))}</b>", styles["VSBody"]),
        ]
        if ch.get("ai_story") or ch.get("ai_notes"):
            block.append(
                Paragraph(
                    f"<b>Cloud Assistant:</b> {_escape(_safe(ch.get('ai_story') or ch.get('ai_notes'), 1200))}",
                    styles["VSBody"],
                )
            )
        block.append(Paragraph("<b>BEFORE</b>", styles["VSSmall"]))
        block.append(Paragraph(_escape(_safe(ch.get("before"), 2500)), styles["VSHint"]))
        block.append(Paragraph("<b>WHAT CHANGED</b>", styles["VSSmall"]))
        block.append(Paragraph(_escape(_safe(ch.get("what_changed"), 1500)), styles["VSBody"]))
        if ch.get("error"):
            block.append(
                Paragraph(
                    f"<font color='#ff3d57'><b>Error:</b> {_escape(_safe(ch.get('error'), 800))}</font>",
                    styles["VSBody"],
                )
            )
        block.append(Paragraph("<b>AFTER</b>", styles["VSSmall"]))
        block.append(Paragraph(_escape(_safe(ch.get("after"), 1500)), styles["VSBody"]))
        cli = ch.get("cli_text") or "\n".join(ch.get("cli_commands") or [])
        if cli:
            block.append(Paragraph("<b>CLI</b>", styles["VSSmall"]))
            block.append(
                Paragraph(
                    f"<font face='Courier' size='7'>{_escape(_safe(cli, 2000))}</font>",
                    styles["VSHint"],
                )
            )
        block.append(Spacer(1, 8))
        story.append(KeepTogether(block))

    full_cli = report.get("cli_script") or ""
    if full_cli:
        story.append(PageBreak())
        story.append(Paragraph("4. Full CLI script (copy / paste)", styles["VSH2"]))
        story.append(
            Paragraph(
                "Run with credentials for the same lab account as your Settings Role ARN.",
                styles["VSHint"],
            )
        )
        # chunk long CLI
        chunk = _safe(full_cli, 12000)
        for part in chunk.split("\n"):
            story.append(
                Paragraph(
                    f"<font face='Courier' size='7'>{_escape(part) if part else '&nbsp;'}</font>",
                    styles["VSHint"],
                )
            )

    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            "VaultScan · Remediation Change Report · Confidential",
            styles["VSCenter"],
        )
    )
    doc.build(story)
    return buf.getvalue()


def export_fix_report_docx(report: dict[str, Any]) -> bytes:
    """Word document version of the fix change report."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading("Remediation Change Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    counts = report.get("counts") or {}
    delta = report.get("score_delta")
    delta_s = "—" if delta is None else f"{delta:+}"

    meta = doc.add_paragraph()
    meta.add_run(
        f"Report: {report.get('report_id')}  ·  Job: {report.get('job_id')}  ·  "
        f"Status: {report.get('job_status')}\n"
        f"Score: {report.get('score_before')} → {report.get('score_after')} (Δ {delta_s})\n"
        f"Applied {counts.get('applied', 0)} · Failed {counts.get('failed', 0)} · "
        f"Skipped {counts.get('skipped', 0)} · Total {counts.get('total', 0)}\n"
        f"Generated {_utc_now()}"
        + (" · Cloud Assistant" if report.get("ai_used") else "")
    ).font.size = Pt(9)

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(_safe(report.get("executive_summary"), 5000))

    recs = report.get("recommendations") or []
    if recs:
        doc.add_heading("2. Recommendations", level=1)
        for r in recs:
            doc.add_paragraph(_safe(r, 800), style="List Bullet")

    doc.add_heading("3. Detailed changes (before → after)", level=1)
    for i, ch in enumerate(report.get("changes") or [], 1):
        doc.add_heading(
            f"{i}. {_safe(ch.get('rule_id'), 40)} — {_safe(ch.get('status'), 24)}",
            level=2,
        )
        p = doc.add_paragraph()
        p.add_run(f"Resource: ").bold = True
        p.add_run(_safe(ch.get("resource"), 120)).font.name = "Courier New"
        if ch.get("title") or ch.get("summary"):
            doc.add_paragraph(_safe(ch.get("title") or ch.get("summary"), 300))
        if ch.get("ai_story") or ch.get("ai_notes"):
            ap = doc.add_paragraph()
            ap.add_run("Cloud Assistant: ").bold = True
            ap.add_run(_safe(ch.get("ai_story") or ch.get("ai_notes"), 1500))

        for label, key in (
            ("BEFORE", "before"),
            ("WHAT CHANGED", "what_changed"),
            ("AFTER", "after"),
        ):
            h = doc.add_paragraph()
            h.add_run(label).bold = True
            doc.add_paragraph(_safe(ch.get(key), 3000))

        if ch.get("error"):
            ep = doc.add_paragraph()
            run = ep.add_run("Error: " + _safe(ch.get("error"), 1000))
            run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)

        cli = ch.get("cli_text") or "\n".join(ch.get("cli_commands") or [])
        if cli:
            cp = doc.add_paragraph()
            cp.add_run("CLI").bold = True
            cli_p = doc.add_paragraph(_safe(cli, 4000))
            for run in cli_p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8)

    full_cli = report.get("cli_script") or ""
    if full_cli:
        doc.add_heading("4. Full CLI script", level=1)
        doc.add_paragraph(
            "Run with credentials for the same lab account as your Settings Role ARN."
        )
        cli_p = doc.add_paragraph(_safe(full_cli, 15000))
        for run in cli_p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    foot = doc.add_paragraph()
    foot.add_run(
        "VaultScan · Remediation Change Report · Confidential"
    ).font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
